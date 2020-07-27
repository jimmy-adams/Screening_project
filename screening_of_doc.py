# -*- coding: utf-8 -*-
"""
Created on Thu Jul 23 16:15:35 2020

@author: Melody
"""
import docx

file_path = "C:/Users/86158/Desktop/项目组成员信息.docx"
file = docx.Document(file_path)
#read the paragraph
print("段落数:"+str(len(file.paragraphs)))#段落数为13，每个回车隔离一段
for para in file.paragraphs:
    print(para.text)
    for i in range(len(para.runs)):
        if "深圳" in para.runs[i].text:
            print("found it")
            para.runs[i].text = para.runs[i].text.replace("深圳", "ShenZhen")

# read the tables
tables = file.tables
print("表格数:"+str(len(file.tables)))#段落数为13，每个回车隔离一段
for table in tables[:]:
    for i, row in enumerate(table.rows[:]):   # 读每行
        row_content = []
        for cell in row.cells[:]:  # 读一行中的所有单元格
            c = cell.text
            row_content.append(c)
        print (row_content) #以列表形式导出每一行数据

# read the header and 
sections = file.sections
